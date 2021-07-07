import docx
from docx.shared import Inches


def add_picture(file_name):
    doc = docx.Document(f"./outputDocs/{file_name}")
    tables = doc.tables
    p = tables[0].rows[0].cells[0].add_paragraph()
    r = p.add_run()
    r.add_picture('user_qr.png', width=Inches(1.2), height=Inches(1.2))
    doc.save(f"./outputDocs/{file_name}")
