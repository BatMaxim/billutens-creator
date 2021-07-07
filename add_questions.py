import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.xmlchemy import OxmlElement
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import _Cell
from docx.shared import Pt

ans_list = ["За", "", "Против", "", "Воздержался"]


def set_cell_border(cell: _Cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def del_line_space(table):
    for row in range(2):
        for col in range(5):
            cell = table.cell(row, col)
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)


def add_table(doc, paragraph):
    table = doc.add_table(rows=2, cols=5)
    for row in range(2):
        for col in range(0, 5, 2):
            cell = table.cell(row, col)
            set_cell_border(cell, top={"sz": 6, "val": "single"},
                            bottom={"sz": 6, "val": "single"},
                            start={"sz": 6, "val": "single"},
                            end={"sz": 6, "val": "single"})
            if row == 0:
                cell.text = ans_list[col]

            cell.paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    del_line_space(table)
    move_table_after(table, paragraph)


def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


def add_questions(ques_list):
    doc = docx.Document('./config/bulletin.docx')

    for paragraph in doc.paragraphs:
        if paragraph.text == "{{ questions }}":
            ques_par = paragraph

    ques_par.clear()
    for i, item in enumerate(ques_list, 1):
        new_p = OxmlElement("w:p")
        ques_par._p.addnext(new_p)
        target_par = Paragraph(new_p, ques_par._parent)
        target_par.add_run(f"{i}. " if i == 1 else f"\n{i}. ")
        run_q = target_par.add_run("Вопрос: ")
        run_q.bold = True
        target_par.add_run(f"{item['question']}\n")
        run_s = target_par.add_run("Решение: ")
        run_s.bold = True
        target_par.add_run(f"{item['solution']}")
        target_par.style.font.name = 'Times New Roman'
        if i != 1:
            add_table(doc, ques_par)
        ques_par = target_par
    add_table(doc, ques_par)
    doc.save('./config/bulletin_with_questions.docx')
